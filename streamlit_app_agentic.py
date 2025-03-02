import streamlit as st
from docx import Document
from openai import OpenAI
import os
import pickle
import requests
import io
import nltk
from sentence_transformers import SentenceTransformer, util
import re
__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
#torch
#pysqlite3-binary
#crewai
#chromadb
#asyncio
from crewai import Agent,Task,Crew,Process,LLM



# from openai import OpenAI
API_KEY=st.secrets.api_key
API_URL = 'https://api.openai.com/v1/chat/completions'

#creating text windows

# def text_break(text):
#     sentences=nltk.sent_tokenize(text)
#     current_len=0
#     current_text=''
#     groups=[]
#     for sentence in sentences:
#         word_len=len(nltk.word_tokenize(sentence))
#         if current_len+word_len<=250:
#             current_len+=word_len
#             current_text+=sentence+' '
#         else:
#             groups.append(current_text.strip())
#             current_text=''
#             current_len=0
#     if current_text:
#         groups.append(current_text.strip())
#     # req_w=round(len(nltk.word_tokenize(text))*1.05)
#     return groups

# Load instructions from a file
def load_instructions(filename):
    """ Load editing instructions from a file. """
    try:
        with open(filename, 'r') as file:
            data = file.read()
            return data
            st.write('file returned')
    except Exception as e:
        st.error('An error occcured:{e}')
        return "Default instructions if file doesn't exist."

def read_docx(file_path):
    """ Extract text from DOCX file. """
    doc = Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    # text=[paragraph.text for paragraph in doc.paragraphs]
    # st.write(text)
    # return text

def tokenize_text(text):
    return set(re.findall(r'\b\w+\b', text.lower()))
    
def process_text_with_api(text, instructions):
    max_t=int(len(nltk.word_tokenize(text))*5.5)
    llm=LLM(model='openai/o1-mini',temperature=0.75,max_tokens=max_t,api_key=API_KEY)
    res_agent=Agent(
        role = "Research Paper Editor",
        goal = """To edit the given research paper strictly as per the provided instructions.""",
        backstory = "You are an expert at editing research papers using the given instructions and have a keen eye for detail and keeping up with the provided guidelines for editing.",
        llm=llm,
        max_iterations=2,
        )

    task = Task(
        description = f"""Edit the research paper as per the provided instructions and striclty follow the guidelines.Do not increase the word count above the original word count.
        The research paper: {text}
        The instructions are: {instructions}""",
        agent=res_agent,
        expected_output = "Edited research paper",
    )

    ppt_crew = Crew(agents=[res_agent], tasks=[task],process=Process.sequential)
    try:
        result=ppt_crew.kickoff()
        return result.raw
    except Exception as e:
        return f"An error occurred: {e}"

def process_document(filename, options,report_features,edits,style):
    """ Read the DOCX file, process the text with loaded instructions and additional features, call the API. """
    text = read_docx(filename)
    # groups = text_break(text)
    instructions = load_instructions(options)
    combined_text = instructions + " " + " ".join([report_features[feature] for feature in edits])
    return process_text_with_api(text,combined_text) 

def create_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    # Save the document to a BytesIO object
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
#################################################################
st.cache_data()
nltk.download('punkt_tab')
model = SentenceTransformer('all-MiniLM-L6-v2') 
# nlp = spacy.load("en_core_web_md")


# Define the path for instruction files
edit_config = {
    "Standard": './standard_new.txt',
    "Developmental": './developmental.txt',
    "ProofReading": './proofreading_new.txt'
}

# Report features for additional processing options
report_features = {
    "aspects of research": """State the novelty, contributions, relevance, findings, practical implications, future scope of research and limitations of this research (30-50 words for each point).""",
    "research_gaps": """Identify the gaps in information presented in the current article in terms of novelty, methodology details, contributions, relevance, findings, practical implications, future scope, and limitations. Check if these aspects of information are explicitly stated in the Abstract (all), Introduction (existing research gap, motivation, rationale for methodology, hypotheses and possible contribution), Methodology (reproducibility), Results (findings), Discussion (data interpretation, contribution, relevance, limitations), and Conclusions section (all).""",
    "synopsis": """Write a synopsis/plain language summary for this study based on the following guidelines: The rationale/motivation/problem statement behind the study and the reason behind the chosen line of investigation. The aim/purpose of the study and the intended outcome. A brief summary of the methodology employed in simple language with the techniques highlighted. The key results obtained from the study and the inferences that can be drawn. The major conclusions drawn from the study. The implications of the study results in the area of research and the industry problem that can be targeted through these results (if applicable).""",
    "highlights": """Outline key highlights. Make five separate points, each for the novelty, contributions, relevance, findings, and practical implications into complete sentences of 10-12 words as Highlights of the research (within 85 characters including spaces).""",
    "title": """Suggest a concise and descriptive title between 10-12 words that incorporates the key content/topics and highlights the main novelty/impact/contribution in the simplest & most comprehensive manner. Max length: 15 words.""",
    "keywords": """Extract significant keywords that represent the current research topic and could be classified as relevant keywords to improve the discoverability of the present topic.""",
    "cover_letter": """Write a cover letter addressed to a journal editor succinctly pitching the main objective, motivation (research gap addressed), methodology, contributions, relevance, key findings, and practical implications of this study."""
}
# Streamlit UI
st.title("Text Editor")

edit_styles=['Standard','Developmental','ProofReading']
style = st.selectbox("Select the type of editing", edit_styles)
options=edit_config[style]

uploaded_file=st.file_uploader("Upload the text document to process.",type=["docx","pdf","csv"],key='orig')
edits = st.multiselect('Select required features:',report_features)
# max_tokens=st.number_input("Insert the maximum number of tokens")

if st.button('Edit Text'):
    response=process_document(uploaded_file,options,report_features,edits,style)
    st.write(response)
    docx_file = create_docx(response)
    before_text=read_docx(uploaded_file)
    # text=read_docx(uploaded_file)
    # before_text="\n".join([t for t in text])
    after_text=response
    embeddings1 = model.encode(before_text, convert_to_tensor=True)
    embeddings2 = model.encode(after_text, convert_to_tensor=True)
    semantic_similarity = util.pytorch_cos_sim(embeddings1, embeddings2).item()
    
    # Calculate word overlap ratio
    before_words = tokenize_text(before_text)
    after_words = tokenize_text(after_text)
    common_words = before_words.intersection(after_words)
    word_overlap_ratio = len(common_words) / len(before_words) if before_words else 0
    
    # Print results
    st.write("Semantic Similarity:", round(semantic_similarity, 2))
    st.write("Word Retention Ratio:", round(word_overlap_ratio * 100, 2), "%")
    st.download_button(
        label="Download as DOCX",
        data=docx_file,
        file_name="output.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
